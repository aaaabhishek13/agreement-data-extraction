"""
Document processing utilities for extracting text from various file formats.
Supports PDF, DOCX, and other document types.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

from PIL import Image
import io

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        self.processors = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_docx,  # Will attempt to process as docx
            '.txt': self._extract_from_txt,
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.SUPPORTED_FORMATS:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            processor = self.processors.get(file_extension)
            if not processor:
                raise ValueError(f"No processor available for {file_extension}")
            
            result = processor(file_path)
            
            # Add metadata
            result.update({
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension,
                'processing_status': 'success'
            })
            
            logger.info(f"Successfully extracted text from {file_path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return {
                'text': '',
                'error': str(e),
                'processing_status': 'failed',
                'file_name': file_path.name if 'file_path' in locals() else 'unknown'
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file with page-wise information."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        text = ""
        page_count = 0
        pages_data = []  # Store page-wise information
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # Clean and format page text
                            cleaned_text = page_text.strip()
                            
                            # Store page information
                            pages_data.append({
                                'page_number': page_num + 1,
                                'text': cleaned_text,
                                'text_length': len(cleaned_text)
                            })
                            
                            # Add to combined text with clear page markers
                            text += f"\n\n=== PAGE {page_num + 1} ===\n"
                            text += cleaned_text
                            text += f"\n=== END PAGE {page_num + 1} ===\n"
                            
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        pages_data.append({
                            'page_number': page_num + 1,
                            'text': '',
                            'error': str(e)
                        })
                        continue
                        
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        
        return {
            'text': text.strip(),
            'page_count': page_count,
            'pages_data': pages_data,
            'extraction_method': 'PyPDF2',
            'supports_citations': True
        }
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file with section information."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX processing")
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs with section context
            paragraphs = []
            sections_data = []
            current_section = 1
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraph_text = para.text.strip()
                    paragraphs.append(paragraph_text)
                    
                    # Create section markers every 10 paragraphs or for headers
                    if (i > 0 and i % 10 == 0) or (para.style and 'heading' in para.style.name.lower()):
                        current_section += 1
                    
                    sections_data.append({
                        'section': current_section,
                        'paragraph': i + 1,
                        'text': paragraph_text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables with section context
            table_data = []
            for table_num, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(" | ".join(row_text))
                if table_text:
                    table_content = "\n".join(table_text)
                    table_data.append(table_content)
                    sections_data.append({
                        'section': f'Table {table_num + 1}',
                        'text': table_content,
                        'type': 'table'
                    })
            
            # Combine all text with section markers
            text = ""
            for i, para in enumerate(paragraphs):
                if i > 0 and i % 10 == 0:
                    text += f"\n\n=== SECTION {(i // 10) + 1} ===\n"
                text += para + "\n"
            
            if table_data:
                text += "\n\n=== TABLES SECTION ===\n"
                text += "\n\n".join(table_data)
            
            return {
                'text': text.strip(),
                'paragraph_count': len(paragraphs),
                'table_count': len(table_data),
                'sections_data': sections_data,
                'extraction_method': 'python-docx',
                'supports_citations': True  # We can provide section/paragraph references
            }
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    # Add simple line-based citations for text files
                    lines = text.split('\n')
                    lines_data = []
                    for i, line in enumerate(lines):
                        if line.strip():
                            lines_data.append({
                                'line_number': i + 1,
                                'text': line.strip()
                            })
                    
                    return {
                        'text': text,
                        'encoding': encoding,
                        'extraction_method': 'plain_text',
                        'lines_data': lines_data,
                        'supports_citations': True  # Line-based citations
                    }
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    def validate_file(self, file_path: str, max_size_mb: int = 16) -> Dict[str, Any]:
        """
        Validate if a file can be processed.
        
        Args:
            file_path: Path to the file
            max_size_mb: Maximum file size in MB
            
        Returns:
            Dict with validation results
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {'valid': False, 'error': 'File does not exist'}
            
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > max_size_mb * 1024 * 1024:
                return {
                    'valid': False, 
                    'error': f'File size ({file_size / (1024*1024):.1f}MB) exceeds limit ({max_size_mb}MB)'
                }
            
            # Check file extension
            file_extension = file_path.suffix.lower()
            if file_extension not in self.SUPPORTED_FORMATS:
                return {
                    'valid': False,
                    'error': f'Unsupported file format: {file_extension}. Supported: {", ".join(self.SUPPORTED_FORMATS)}'
                }
            
            return {
                'valid': True,
                'file_size': file_size,
                'file_extension': file_extension,
                'file_name': file_path.name
            }
            
        except Exception as e:
            return {'valid': False, 'error': f'Validation error: {str(e)}'}
